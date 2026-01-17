"""
Generate IEEE format Word document report for Wi-Fi CSI localization project.
"""

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width

def add_two_column_section(doc):
    """Configure document for two-column layout."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1)

    # Set two columns
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '360')  # 0.25 inch between columns
    sectPr.append(cols)

def add_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'

def add_authors(doc, authors, affiliation):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authors)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(affiliation)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.space_before = Pt(12)
    p.space_after = Pt(6)

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.space_before = Pt(6)
    p.space_after = Pt(3)

def add_body_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.space_after = Pt(3)

def add_body_text_no_indent(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(3)

def add_bullet_point(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def add_figure_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[INSERT FIGURE HERE]")
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(6)

def add_table(doc, headers, data, caption):
    # Add caption above table
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    p.space_before = Pt(6)

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

    # Add data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # Space after table

def main():
    doc = Document()

    # Configure two-column layout
    add_two_column_section(doc)

    # Title
    add_title(doc, "Wi-Fi Device Localization Using Deep Neural Networks and CSI Feature Engineering")

    # Authors
    add_authors(doc, "Author Name 1, Author Name 2", "Department, University")

    doc.add_paragraph()  # Spacing

    # Abstract
    add_section_heading(doc, "Abstract")

    abstract_text = (
        "This paper presents a deep learning approach for Wi-Fi device localization using "
        "Channel State Information (CSI). The task is formulated as a 10-class classification "
        "problem where the goal is to predict the position of a Wi-Fi device based on "
        "260-dimensional CSI measurements. We initially explored Random Forest classifiers "
        "but found that neural networks better capture the complex patterns in high-dimensional "
        "CSI data. Our approach combines domain-specific feature engineering with residual "
        "neural networks. The feature engineering pipeline transforms raw I/Q values into "
        "amplitude and phase representations, removes inactive guard band subcarriers, and "
        "extracts cross-antenna features that capture angle-of-arrival information. We "
        "conducted a hyperparameter search over 43 configurations and built a 12-model "
        "ensemble using diversity-based selection. The final ensemble achieves 96.59% "
        "validation accuracy, representing a 0.88% improvement over our baseline feedforward network."
    )
    add_body_text_no_indent(doc, abstract_text)

    # I. Introduction
    add_section_heading(doc, "I. INTRODUCTION")

    add_body_text_no_indent(doc,
        "Channel State Information (CSI) describes how wireless signals are affected by "
        "the environment before reaching a receiver. Modern Wi-Fi systems use Orthogonal "
        "Frequency Division Multiplexing (OFDM), which splits transmissions across multiple "
        "closely spaced subcarriers. As signals travel through space, they experience "
        "reflections, scattering, and fading. CSI captures these effects by measuring "
        "amplitude and phase for each subcarrier."
    )

    add_body_text(doc,
        "CSI-based sensing has become increasingly important in recent years. The high "
        "dimensionality of CSI data makes it well suited for machine learning approaches, "
        "which can learn complex patterns that traditional analytical models struggle to "
        "capture. Applications include motion detection, presence sensing, occupancy "
        "tracking, and device localization."
    )

    add_body_text(doc,
        "This paper addresses the problem of Wi-Fi device localization. Given CSI "
        "measurements from a tracking unit, we aim to classify the position of a target "
        "Wi-Fi device into one of ten predefined classes. The dataset contains 12,888 "
        "training samples with 260 features each. The features consist of 5 metadata "
        "values (timestamp, sequence control, angle of arrival, and two RSSI measurements) "
        "plus 255 CSI values representing I/Q pairs for 64 subcarriers across 2 antennas."
    )

    add_body_text(doc, "Our main contributions are:")

    add_bullet_point(doc, "A feature engineering pipeline that transforms raw CSI data into more informative representations")
    add_bullet_point(doc, "A residual neural network architecture designed for CSI classification")
    add_bullet_point(doc, "A diversity-optimized ensemble that combines multiple trained models")
    add_bullet_point(doc, "Comprehensive hyperparameter optimization across 43 configurations")

    # II. Feature Analysis
    add_section_heading(doc, "II. FEATURE ANALYSIS")

    add_subsection_heading(doc, "A. Raw CSI Structure")

    add_body_text_no_indent(doc,
        "The raw CSI data consists of 255 values representing measurements from 64 OFDM "
        "subcarriers across 2 receiver antennas. Each subcarrier measurement includes "
        "in-phase (I) and quadrature (Q) components. The 64 subcarriers correspond to a "
        "20 MHz Wi-Fi channel."
    )

    add_body_text(doc,
        "Not all subcarriers carry useful information. Wi-Fi channels include guard bands "
        "at the edges and null subcarriers to prevent interference between adjacent channels. "
        "In our dataset, 12 subcarriers are inactive: the DC subcarrier (index 0) and guard "
        "band subcarriers (indices 27-37)."
    )

    add_subsection_heading(doc, "B. Feature Engineering Pipeline")

    add_figure_placeholder(doc, "Fig. 1. Feature engineering pipeline transforming raw CSI to 335 engineered features.")

    add_body_text_no_indent(doc,
        "Our feature engineering pipeline transforms 260 raw features into 335 engineered "
        "features through the following steps:"
    )

    add_body_text(doc,
        "1) Amplitude and Phase Extraction: Raw I/Q values are converted to polar "
        "representation. For each subcarrier, amplitude is computed as sqrt(I^2 + Q^2) "
        "and phase as arctan2(Q, I). This representation separates signal strength from "
        "timing information."
    )

    add_body_text(doc,
        "2) Inactive Subcarrier Removal: We remove the 12 inactive subcarriers (indices "
        "0, 27-37) to reduce noise from null and pilot carriers. This leaves 52 active "
        "subcarriers per antenna."
    )

    add_body_text(doc,
        "3) Cross-Antenna Features: We compute amplitude difference and normalized phase "
        "difference between the two antennas for each active subcarrier. These features "
        "capture angle-of-arrival information that helps distinguish device positions."
    )

    add_body_text(doc,
        "4) Band Statistics: We divide subcarriers into low, mid, and high frequency bands "
        "and compute aggregate statistics (mean, standard deviation, maximum) for each band "
        "and antenna. This captures frequency-dependent propagation characteristics."
    )

    # III. Classification Methods
    add_section_heading(doc, "III. CLASSIFICATION METHODS")

    add_subsection_heading(doc, "A. Initial Approach: Random Forest")

    add_body_text_no_indent(doc,
        "We initially experimented with Random Forest classifiers. Random Forests are a "
        "natural choice for tabular data and provide good baseline performance without "
        "extensive tuning. However, we found that Random Forests struggled to capture the "
        "complex spatial relationships in high-dimensional CSI data. The I/Q values from "
        "different subcarriers and antennas have intricate correlations that tree-based "
        "methods do not model effectively."
    )

    add_subsection_heading(doc, "B. Neural Network Approach")

    add_body_text_no_indent(doc,
        "We switched to neural networks, which can learn arbitrary nonlinear relationships "
        "between features. Neural networks are particularly well suited for CSI data because "
        "they can jointly process all subcarrier measurements and learn representations that "
        "capture spatial patterns."
    )

    add_subsection_heading(doc, "C. Baseline Architecture (v3)")

    add_figure_placeholder(doc, "Fig. 2. Neural network architectures: v3 baseline (left) and v4 ResNet (right).")

    add_body_text_no_indent(doc,
        "Our baseline model is a 4-layer feedforward network with the following structure: "
        "Input (260) -> Linear(512) -> BatchNorm -> ReLU -> Dropout(0.3) -> Linear(256) -> "
        "BatchNorm -> ReLU -> Dropout(0.3) -> Linear(128) -> BatchNorm -> ReLU -> Dropout(0.2) "
        "-> Linear(64) -> BatchNorm -> ReLU -> Dropout(0.2) -> Output(10)."
    )

    add_body_text(doc,
        "Batch normalization stabilizes training and allows higher learning rates. Dropout "
        "provides regularization to prevent overfitting. The network outputs raw logits, "
        "and softmax is applied internally by the cross-entropy loss function."
    )

    add_subsection_heading(doc, "D. Improved Architecture (v4)")

    add_body_text_no_indent(doc,
        "The improved v4 model uses residual blocks with skip connections. The architecture "
        "is: Input (335) -> Linear(1024) -> BatchNorm -> ReLU -> Dropout(0.45) -> ResBlock "
        "(1024->512) -> ResBlock (512->256) -> ResBlock (256->128) -> ResBlock (128->64) -> "
        "Linear(64) -> BatchNorm -> ReLU -> Dropout(0.2) -> Output(10)."
    )

    add_body_text(doc,
        "Each residual block consists of two linear layers with batch normalization, and a "
        "skip connection that adds the input to the output. Skip connections help gradient "
        "flow during training and allow the network to learn residual mappings."
    )

    add_subsection_heading(doc, "E. Ensemble Strategy")

    add_body_text_no_indent(doc,
        "Our final model is an ensemble of 12 neural networks. Individual predictions are "
        "averaged (soft voting) to produce the final class probabilities. The ensemble "
        "members are selected using a diversity-based greedy algorithm: (1) Train 43 models "
        "with different hyperparameter configurations, (2) Select the top 20 models by "
        "validation accuracy, (3) Start with the best model, (4) Iteratively add models "
        "that maximize 0.7*accuracy + 0.3*diversity, where diversity is measured by "
        "prediction disagreement rate."
    )

    # IV. Evaluation Metrics
    add_section_heading(doc, "IV. EVALUATION METRICS")

    add_body_text_no_indent(doc,
        "The primary evaluation metric is classification accuracy, which measures the "
        "fraction of correctly classified samples. We also report precision, recall, "
        "and F1-score to understand per-class performance. Confusion matrix analysis "
        "reveals which position classes are commonly confused, providing insight into "
        "the spatial relationships between positions."
    )

    # V. Experiments
    add_section_heading(doc, "V. EXPERIMENTS")

    add_subsection_heading(doc, "A. Data Split and Preprocessing")

    add_body_text_no_indent(doc,
        "We use an 85/15 stratified train/validation split. Stratification ensures each "
        "split has the same class distribution as the original dataset. All features are "
        "standardized (zero mean, unit variance) using statistics computed on the training set."
    )

    add_subsection_heading(doc, "B. Hyperparameter Search")

    add_body_text_no_indent(doc,
        "We searched 43 hyperparameter configurations, varying learning rate (0.0005 to "
        "0.0012), weight decay (5e-6 to 5e-5), dropout base (0.35 to 0.45), label smoothing "
        "(0.05 to 0.15), and mixup alpha (0.15 to 0.25). Table I shows the top configurations."
    )

    # Table I: Hyperparameter configurations
    headers = ["Rank", "LR", "Dropout", "Label Sm.", "Mixup", "Val Acc"]
    data = [
        ["1", "0.001", "0.45", "0.1", "0.20", "96.48%"],
        ["2", "0.001", "0.45", "0.1", "0.25", "96.38%"],
        ["3", "0.001", "0.35", "0.05", "0.20", "96.28%"],
        ["4", "0.001", "0.45", "0.05", "0.20", "96.23%"],
        ["5", "0.001", "0.35", "0.15", "0.20", "96.23%"],
    ]
    add_table(doc, headers, data, "TABLE I: TOP HYPERPARAMETER CONFIGURATIONS")

    add_subsection_heading(doc, "C. Training Configuration")

    add_body_text_no_indent(doc,
        "All models are trained with AdamW optimizer, batch size 256, maximum 1000 epochs, "
        "early stopping patience of 30-40 epochs, 10-epoch learning rate warmup followed by "
        "cosine annealing, and gradient clipping with max norm 1.0."
    )

    add_subsection_heading(doc, "D. Training Enhancements")

    add_body_text_no_indent(doc,
        "Mixup Data Augmentation: During training, we linearly interpolate between random "
        "pairs of samples and their labels. Given samples (x_i, y_i) and (x_j, y_j), we "
        "create mixed samples where lambda is sampled from Beta(alpha, alpha). This smooths "
        "the decision boundary and improves generalization."
    )

    add_body_text(doc,
        "Label Smoothing: Instead of hard targets (0 or 1), we use soft targets that "
        "distribute a small amount of probability mass to incorrect classes. This prevents "
        "overconfident predictions and improves calibration."
    )

    # VI. Results
    add_section_heading(doc, "VI. RESULTS")

    add_figure_placeholder(doc, "Fig. 3. Training and validation accuracy curves over epochs.")

    # Table II: Model comparison
    headers2 = ["Model", "Val. Accuracy", "Improvement"]
    data2 = [
        ["v3 Baseline", "95.71%", "-"],
        ["v4 Single (best)", "96.48%", "+0.77%"],
        ["v4 Ensemble (12)", "96.59%", "+0.88%"],
    ]
    add_table(doc, headers2, data2, "TABLE II: MODEL COMPARISON RESULTS")

    add_subsection_heading(doc, "A. Baseline Performance")

    add_body_text_no_indent(doc,
        "The v3 baseline model achieves 95.71% validation accuracy. This model uses raw "
        "features without engineering and a simple feedforward architecture. Training "
        "converges within approximately 200 epochs with early stopping."
    )

    add_subsection_heading(doc, "B. Feature Engineering Impact")

    add_body_text_no_indent(doc,
        "Adding feature engineering improves single-model accuracy from 95.71% to 96.48%. "
        "The amplitude/phase representation and cross-antenna features provide the most "
        "benefit, as they directly encode information relevant to device localization."
    )

    add_subsection_heading(doc, "C. Ensemble Performance")

    add_body_text_no_indent(doc,
        "The 12-model diversity-optimized ensemble achieves 96.59% validation accuracy. "
        "Compared to simple top-N averaging, diversity-based selection produces better "
        "results with fewer models. The ensemble corrects errors made by individual models "
        "while introducing few new errors."
    )

    add_subsection_heading(doc, "D. Best Hyperparameters")

    add_body_text_no_indent(doc,
        "The best single model uses: learning rate 0.001, weight decay 5e-6, dropout base "
        "0.45, label smoothing 0.1, and mixup alpha 0.2. Higher dropout combined with label "
        "smoothing and mixup provides strong regularization without degrading training."
    )

    # VII. Conclusions
    add_section_heading(doc, "VII. CONCLUSIONS")

    add_body_text_no_indent(doc,
        "We presented a deep learning system for Wi-Fi device localization using CSI data. "
        "Our main findings are: (1) Feature engineering significantly improves performance "
        "by converting I/Q values to amplitude and phase and extracting cross-antenna "
        "features; (2) Residual networks outperform feedforward networks due to better "
        "gradient flow; (3) Ensemble diversity matters more than simply averaging top "
        "performers; (4) Regularization through dropout, label smoothing, and mixup is "
        "crucial for preventing overfitting."
    )

    add_body_text(doc,
        "Our final ensemble achieves 96.59% validation accuracy on the 10-class "
        "localization task. Future work could explore 1D convolutional networks, attention "
        "mechanisms, or CSI-specific data augmentation techniques."
    )

    # References
    add_section_heading(doc, "REFERENCES")

    refs = [
        "[1] A. Paszke et al., \"PyTorch: An Imperative Style, High-Performance Deep Learning Library,\" in Advances in Neural Information Processing Systems 32, 2019.",
        "[2] K. He, X. Zhang, S. Ren, and J. Sun, \"Deep Residual Learning for Image Recognition,\" in Proc. IEEE CVPR, 2016.",
        "[3] H. Zhang, M. Cisse, Y. N. Dauphin, and D. Lopez-Paz, \"mixup: Beyond Empirical Risk Minimization,\" in ICLR, 2018.",
        "[4] C. Szegedy et al., \"Rethinking the Inception Architecture for Computer Vision,\" in Proc. IEEE CVPR, 2016.",
        "[5] IEEE 802.11bf Task Group, \"WLAN Sensing,\" IEEE Standards Association, 2024.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        p.space_after = Pt(2)

    # Save document
    doc.save('IEEE_Report_WiFi_Localization.docx')
    print("Report saved to 'IEEE_Report_WiFi_Localization.docx'")

if __name__ == '__main__':
    main()
